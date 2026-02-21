import os
import fitz  # PyMuPDF
import docx
import markdown
from PySide6.QtGui import QTextDocument
from PySide6.QtWidgets import QTextEdit

class FileManager:
    @staticmethod
    def read_file(file_path: str) -> dict:
        """
        Reads a file and returns its content and type.
        Returns a dict: {'content': str, 'type': str, 'format': str}
        'format' can be 'html' or 'plain'.
        """
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        if ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return {'content': f.read(), 'type': 'txt', 'format': 'plain'}
        
        elif ext == '.md':
            with open(file_path, 'r', encoding='utf-8') as f:
                md_text = f.read()
                html_content = markdown.markdown(md_text)
                return {'content': html_content, 'type': 'md', 'format': 'html'}

        elif ext == '.pdf':
            text = ""
            try:
                doc = fitz.open(file_path)
                for page in doc:
                    text += page.get_text()
                return {'content': text, 'type': 'pdf', 'format': 'plain'}
            except Exception as e:
                return {'content': f"Error reading PDF: {str(e)}", 'type': 'pdf', 'format': 'plain'}

        elif ext == '.docx':
            try:
                doc = docx.Document(file_path)
                # Simple extraction, preserving paragraphs but losing complex formatting
                # For better HTML conversion, a dedicated library like mammoth is often used,
                # but we'll stick to basic paragraph extraction for now or basic HTML construction.
                html_content = ""
                for para in doc.paragraphs:
                    html_content += f"<p>{para.text}</p>"
                return {'content': html_content, 'type': 'docx', 'format': 'html'}
            except Exception as e:
                return {'content': f"Error reading DOCX: {str(e)}", 'type': 'docx', 'format': 'plain'}

        return {'content': '', 'type': 'unknown', 'format': 'plain'}

    @staticmethod
    def save_file(file_path: str, content_html: str, content_plain: str):
        """
        Saves content to a file.
        """
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        if ext == '.txt':
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content_plain)
        
        elif ext == '.md':
            # Converting HTML back to MD is complex without a library like markdownify.
            # For now, we save plain text or basic conversion.
            # Ideally, we should use a library or just save the plain text from the editor.
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content_plain) 

        elif ext == '.docx':
             # Writing DOCX from HTML is hard without pypandoc or similar.
             # We will create a basic DOCX with paragraphs from plain text for this prototype.
             doc = docx.Document()
             for line in content_plain.split(''):
                 doc.add_paragraph(line)
             doc.save(file_path)

        elif ext == '.pdf':
            # PDF export should be handled by the UI (QPrinter) for WYSIWYG.
            # This method might not be used for PDF if we print directly from UI.
            pass
